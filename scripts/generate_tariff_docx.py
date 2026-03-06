"""
Generate branded 7Square TPA Tariff Book Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# === Brand colors ===
BRAND_PURPLE = RGBColor(0x1E, 0x0A, 0x3C)  # Dark purple from logo
BRAND_GOLD = RGBColor(0xC5, 0xA4, 0x4E)    # Gold accent from logo
BRAND_LIGHT_PURPLE = RGBColor(0x3D, 0x1D, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)

BASE = r"c:\Users\ernes\Documents\Projects\tpa_sea_ports"
LOGO = os.path.join(BASE, "The _7Square full color logo.png")
OUTPUT = os.path.join(BASE, "7Square_TPA_Tariff_Book_Feb_2024.docx")

doc = Document()

# === Page setup ===
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = DARK_GRAY

# === Helper functions ===

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_branded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1E0A3C')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            # Right-align numeric columns (typically last cols)
            if c_idx > 0 and any(ch.isdigit() for ch in str(val)) and '%' not in str(val):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F0FF')

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    # Reduce cell padding
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

    return table

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_PURPLE
        run.font.name = 'Calibri'
    return h

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = MID_GRAY
    run.font.name = 'Calibri'
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'
    return p

def add_gold_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run('_' * 85)
    run.font.color.rgb = BRAND_GOLD
    run.font.size = Pt(6)

def add_page_break(doc):
    doc.add_page_break()

# ======================================================================
# COVER PAGE
# ======================================================================
for _ in range(4):
    doc.add_paragraph()

# Logo centered
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(LOGO, width=Inches(3.5))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TPA SEA PORTS')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = BRAND_PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TARIFF BOOK')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BRAND_GOLD
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('February 2024 Edition')
run.font.size = Pt(14)
run.font.color.rgb = MID_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Structured Reference for 7Square ICD Operations Platform')
run.font.size = Pt(11)
run.font.color.rgb = BRAND_LIGHT_PURPLE
run.font.name = 'Calibri'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('www.7squareinc.com  |  info@7squareinc.com')
run.font.size = Pt(9)
run.font.color.rgb = MID_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Source: Tanzania Ports Authority | Section 71(1)(a) Ports Act Cap. 166')
run.font.size = Pt(8)
run.font.color.rgb = MID_GRAY

# ======================================================================
# HEADER/FOOTER for subsequent pages
# ======================================================================
section = doc.sections[0]
header = section.header
hp = header.paragraphs[0]
hr = hp.add_run()
hr.add_picture(LOGO, width=Inches(1.2))
hp.add_run('    ')
hr2 = hp.add_run('TPA Sea Ports Tariff Book — February 2024')
hr2.font.size = Pt(8)
hr2.font.color.rgb = MID_GRAY
hr2.font.name = 'Calibri'

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('7Square Inc.  |  Confidential  |  ')
fr.font.size = Pt(7)
fr.font.color.rgb = MID_GRAY
# Page number
fld_xml = (
    '<w:fldSimple {} w:instr=" PAGE "/>'.format(nsdecls('w'))
)
fp._p.append(parse_xml(fld_xml))

add_page_break(doc)

# ======================================================================
# TABLE OF CONTENTS
# ======================================================================
add_heading_styled(doc, 'Table of Contents', 1)
add_gold_divider(doc)

toc_items = [
    ('Part I', 'Application and Interpretations'),
    ('Part II', 'Applicable Principles'),
    ('Part III', 'Fees, Dues and Charges'),
    ('Clause 1', 'Pilotage Fees'),
    ('Clause 2', 'Port Dues'),
    ('Clause 3', 'Navigational Dues'),
    ('Clause 4', 'Dockage and Buoyage'),
    ('Clause 5', 'Tug Services'),
    ('Clause 6', 'Hire of Lighters and Pontoons'),
    ('Clause 7', 'Mooring and Unmooring Services'),
    ('Clause 8', 'Supply of Fresh Water to Vessels'),
    ('Clause 9', 'Garbage Disposal'),
    ('Clause 11', 'Hire of Staff and Labour'),
    ('Clause 12', 'Hire of Equipment'),
    ('Clause 13', 'Port Labour Kept or Remaining Idle'),
    ('Clause 14', 'Stevedoring'),
    ('Clause 15', 'Movement of Cargo in or from a Vessel'),
    ('Clause 16', 'Laid-Up Ships'),
    ('Clause 17', 'Slipping/Unslipping & Hire of Slipways'),
    ('Clause 18', 'Other Charges and Fees'),
    ('Clause 22', 'Private Mooring Buoys'),
    ('Clause 23', 'Amending/Cancelling Orders or Invoices'),
    ('Clause 27', 'Wharfage'),
    ('Clause 28', 'Wayleave Dues'),
    ('Clause 29', 'Shorehandling'),
    ('Clause 30', 'Heavy Lifts'),
    ('Clause 31', 'Removal Charges'),
    ('Clause 32', 'Storage'),
    ('Clause 33', 'Coastwise Cargo'),
    ('Clause 34', 'Import and Export of Livestock'),
    ('Clause 35', 'Special Rates'),
    ('Clause 36', 'Container Handling — DCT Berths 8-11'),
    ('Clause 37', 'Container Handling — Berths 0-7'),
    ('Clause 38', 'Roll On-Roll Off Operations'),
    ('Clause 39', 'Other Container Service Charges'),
    ('Clause 40', 'Handling Charges for Bulk Oils'),
    ('Clause 41', 'Grain Terminal Services'),
    ('Clause 42', 'Special Rate for ICDs / Dry Ports'),
    ('Part IV', 'Clause 43: Miscellaneous Provisions and Charges'),
    ('Appendix A', 'Free Storage Periods Summary'),
    ('Appendix B', 'Global Surcharge Rules Summary'),
]

for ref, title in toc_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{ref}:  ')
    run1.bold = True
    run1.font.size = Pt(9.5)
    run1.font.color.rgb = BRAND_PURPLE
    run2 = p.add_run(title)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(2)

add_page_break(doc)

# ======================================================================
# PART I: APPLICATION AND INTERPRETATIONS
# ======================================================================
add_heading_styled(doc, 'PART I: APPLICATION AND INTERPRETATIONS', 1)
add_gold_divider(doc)

add_body(doc, 'This Tariff Book applies to all sea ports set out in Part I of the Second Schedule to the Ports Act Cap. 166. All charges are denominated in USD. The Authority may allow payment in Tanzanian Shillings at the exchange rate determined by the Authority.')

add_heading_styled(doc, 'Abbreviations', 2)
add_branded_table(doc,
    ['Abbreviation', 'Meaning'],
    [
        ['CIF', 'Cost Insurance and Freight'],
        ['DWT', 'Deadweight Tonne (1,000 kg)'],
        ['FCL', 'Full Container Load'],
        ['GRT', 'Gross Registered Tonnage'],
        ['HTN', 'Harbour Tonne (1,000 kg or 1 CBM, whichever yields higher charge)'],
        ['ICD', 'Inland Container Depot'],
        ['IMDG', 'International Maritime Dangerous Goods'],
        ['KOJ', 'Kurasini Oil Jetty'],
        ['LCL', 'Less Container Load'],
        ['LOA', 'Length Overall'],
        ['SBM', 'Single Buoy Mooring'],
        ['TEU', 'Twenty-foot Equivalent Unit'],
        ['VGM', 'Verified Gross Mass'],
    ],
    col_widths=[1.5, 5.0]
)

add_heading_styled(doc, 'Key Definitions', 2)
defs = [
    ('Coastwise Cargo', 'Cargo carried between ports along Tanzania coastal waters.'),
    ('Dangerous Cargo', 'Substances listed in the IMDG Code (IMO). Must be declared by class. Trade names not acceptable. Attracts 10% surcharge on stevedoring/shorehandling, 20% on storage.'),
    ('Difficult Cargo', 'Charcoal, sulphur, wet hides, pig iron, oil cake, rock sulphate, loose timber, loose scrap, etc.'),
    ('Direct Delivery', 'Cargo discharged from ship directly onto consignee\'s truck without port storage.'),
    ('Domestic', 'Goods produced, permanently entered, or manufactured within Tanzania.'),
    ('Harbour Tonne (HTN)', '1,000 kg or 1 cubic metre, whichever yields the higher charge.'),
    ('Shorehandling', 'Handling/transfer of cargo between quay and transit sheds/warehouses/stacking yards.'),
    ('Stevedoring', 'Transfer of cargo within vessel and/or between vessel and quay or next transport mode.'),
    ('Stuffing/Stripping', 'Loading cargo into / unloading cargo from a container.'),
    ('Valuable Cargo', 'Bullion, precious stones, ivory, essential oils, goldware, platinum, silks, etc.'),
    ('Wharfage', 'Charge on all cargo (incl. empty containers) passing over TPA quays/wharves/jetties/buoys.'),
]
for term, definition in defs:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{term}: ')
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = BRAND_PURPLE
    r2 = p.add_run(definition)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(3)

add_page_break(doc)

# ======================================================================
# PART II: APPLICABLE PRINCIPLES
# ======================================================================
add_heading_styled(doc, 'PART II: APPLICABLE PRINCIPLES', 1)
add_gold_divider(doc)

add_heading_styled(doc, 'Working Hours', 2)
add_branded_table(doc,
    ['Service', 'Regular Hours (Mon-Fri)', 'Overtime'],
    [
        ['Receipt of Import/Export Documents', '0700-2300', 'Sat/Sun/Public Holidays'],
        ['Receipt and Delivery of Cargo', '0700-0700 (24hr)', '—'],
        ['Shorehandling/Stevedoring 1st Shift', '0700-1500', '—'],
        ['Shorehandling/Stevedoring 2nd Shift', '1500-2300', '—'],
        ['Shorehandling/Stevedoring 3rd Shift', '2300-0700', 'Always overtime'],
        ['Conservancy & Water Supply', '0700-1200 & 1300-1600', '—'],
        ['Weekends & Public Holidays', 'All overtime', 'All overtime'],
    ]
)

add_note(doc, 'Invoice Cut-off: Customers can pay and clear cargo on the same day as invoice printed, or within 24 hours, without additional storage charges.')

add_heading_styled(doc, 'Global Surcharge Rules', 2)
add_branded_table(doc,
    ['Surcharge Type', 'Rate', 'Applies To'],
    [
        ['DG — Stevedoring', '+10%', 'Clauses 14, 36, 37, 38'],
        ['DG — Shorehandling', '+10%', 'Clause 29'],
        ['DG — Storage', '+20%', 'Clause 32 (after 24hr free)'],
        ['Over-dimension Containers', '+30%', 'Clauses 29, 36, 37, 38'],
        ['Cold Storage Cargo — Handling', '+30%', 'Clause 29'],
        ['Overtime Gang Charge', '+$500/gang/shift', 'Clauses 14, 36, 37, 38'],
    ]
)

add_heading_styled(doc, 'Container Rules', 2)
add_body(doc, 'Reefer Containers: Charged per Clause 39. Container Status: Must be declared in manifest; undeclared containers attract storage from landing. VGM (SOLAS): Shipper must provide verified gross mass before loading. Heavy lift charges do not apply to containers. Stuffing deadline: export cargo 4 working days before vessel works; empty containers 5 days before.')

add_page_break(doc)

# ======================================================================
# PART III: FEES, DUES AND CHARGES
# ======================================================================
add_heading_styled(doc, 'PART III: FEES, DUES AND CHARGES', 1)
add_gold_divider(doc)

# --- CLAUSE 1 ---
add_heading_styled(doc, 'Clause 1: Pilotage Fees', 2)
add_body(doc, 'Rate per 100 GRT or part thereof, per operation (USD):')
add_branded_table(doc,
    ['#', 'Service', 'Deep Sea', 'Coaster'],
    [
        ['1(a)', 'Entering or leaving port', '5.50', '1.10'],
        ['1(b)', 'Internal movements', '5.50', '1.10'],
        ['1(c)', 'Dead ship movements', '15.00', '3.10'],
        ['1(d)', 'Movements between adjacent berths', '2.80', '0.60'],
        ['1(e)', 'Minimum charge per vessel per service', '150.00', '33.80'],
    ],
    col_widths=[0.5, 3.0, 1.2, 1.2]
)
add_body(doc, 'Pilotage Detention Fees:')
add_branded_table(doc,
    ['Condition', 'Deep Sea (USD)', 'Coaster (USD)'],
    [
        ['First 30 minutes', 'Free', 'Free'],
        ['Thereafter per minute (min $100/$22.50)', '4.20', '0.80'],
    ]
)
add_body(doc, 'Cancellation Fees:')
add_branded_table(doc,
    ['Condition', 'Deep Sea (USD)', 'Coaster (USD)'],
    [
        ['30+ min before service required', 'Free', 'Free'],
        ['Within 30 min of service required', '115.83 (fixed)', '2.20/100 GRT'],
        ['Pilot proceeds, ship doesn\'t pick up', '2x rate 1(a) or 1(b)', '2x rate'],
    ]
)

# --- CLAUSE 2 ---
add_heading_styled(doc, 'Clause 2: Port Dues', 2)
add_body(doc, 'Rate per 100 GRT per call or part thereof (USD):')
add_branded_table(doc,
    ['#', 'Period', 'Deep Sea', 'Coaster'],
    [
        ['1', 'First 5 days or part thereof', '13.40', '2.50'],
        ['2', 'Each successive 5 days', '8.10', '1.50'],
        ['3', 'Bunkering/water/stores', 'Half rates', 'Half rates'],
        ['4(a)', 'Stress/bad weather (first 48 hrs)', 'Free', 'Free'],
        ['4(b)', 'Stress vessels beyond 48 hrs', 'Half rates', 'Half rates'],
        ['5', 'Slipping or dry docking', 'Half rates', 'Half rates'],
    ],
    col_widths=[0.5, 3.0, 1.2, 1.2]
)
add_body(doc, 'Pleasure crafts over 30 GRT (rate per 1 GRT):')
add_branded_table(doc,
    ['Period', 'Deep Sea', 'Coaster'],
    [
        ['First 5 days', '5.40', '1.00'],
        ['Thereafter per 5 days', '3.40', '0.60'],
    ]
)

# --- CLAUSE 3 ---
add_heading_styled(doc, 'Clause 3: Navigational Dues', 2)
add_body(doc, 'Rate per 100 GRT or part thereof, per call (USD):')
add_branded_table(doc,
    ['Service', 'Deep Sea', 'Coaster'],
    [
        ['Combined navigational dues', '6.00', '1.20'],
        ['Minimum charge per call', '26.90', '5.10'],
    ]
)
add_note(doc, 'Annual prepayment at 10x call rate. Charged once per turnaround voyage.')

# --- CLAUSE 4 ---
add_heading_styled(doc, 'Clause 4: Dockage and Buoyage', 2)
add_body(doc, 'Rate per 100 GRT per hour or part thereof (USD):')
add_branded_table(doc,
    ['#', 'Location', 'Deep Sea', 'Coaster'],
    [
        ['1', 'Quays, wharves, jetties', '0.50', '0.10'],
        ['2', 'Moored at buoys', '0.30', '0.10'],
        ['3', 'Double banked', '0.50', '0.10'],
        ['4', 'Bulk Oil Jetties', '0.50', '0.10'],
        ['5', 'Single Mooring Points', '0.50', '0.10'],
        ['6', 'Dhow/Lighter Wharf', '0.30', '0.10'],
        ['7', 'RORO stern ramp', '0.30', '0.10'],
        ['8', 'Outer anchorage (non-documented)', '2.10', 'N/A'],
        ['9', 'Coaster at buoy canvassing', '25% of normal', '—'],
    ],
    col_widths=[0.4, 3.0, 1.2, 1.2]
)

# --- CLAUSE 5 ---
add_heading_styled(doc, 'Clause 5: Tug Services', 2)
add_body(doc, 'Vessels under own steam — Rate per 100 GRT per tug per operation (USD):')
add_branded_table(doc,
    ['#', 'Service', 'Deep Sea', 'Coaster'],
    [
        ['3(a)', 'Berthing or unberthing', '14.00', '5.40'],
        ['3(b)', 'Assisting turning / clearing hawser', '9.40', '2.70'],
        ['3(c)', 'Moving within port limits', '14.00', '5.40'],
        ['3(d)', 'Moving from/to outside port limits', '20.00', '6.80'],
        ['3(e)', 'Beyond 2 nautical miles', 'On application', '—'],
        ['3(f)', 'Towage lighters/pontoons/small crafts', '8.00', '2.70'],
    ],
    col_widths=[0.5, 3.0, 1.2, 1.2]
)
add_note(doc, 'Not under own steam: Double the above rates. Tugs idle: $200/tug/hour (deep sea).')

# --- CLAUSE 6 ---
add_heading_styled(doc, 'Clause 6: Hire of Lighters and Pontoons', 2)
add_body(doc, 'Rate per ton of lighter/pontoon capacity per 12 hours (USD):')
add_branded_table(doc,
    ['#', 'Service', 'Deep Sea', 'Coaster'],
    [
        ['2(a)', 'Loading/discharging in steam', '1.30', '0.40'],
        ['2(b)', 'Alongside quays/wharves/jetties', '1.30', '0.40'],
        ['2(c)', 'As fenders', '1.30', '0.40'],
        ['2(d)', 'Transshipment ship-to-ship', '1.30', '0.60'],
        ['2(e)', 'Explosives/DG cargo', 'Treble rates', 'Treble rates'],
        ['3(a)', 'Other purposes', '1.60', '0.60'],
        ['4', 'Minimum charge', '200.00', '56.40'],
    ],
    col_widths=[0.5, 3.0, 1.2, 1.2]
)

add_page_break(doc)

# --- CLAUSE 7 ---
add_heading_styled(doc, 'Clause 7: Mooring and Unmooring Services', 2)
add_body(doc, 'Rate per 100 GRT per operation (USD):')
add_branded_table(doc,
    ['#', 'Service', 'DS Normal', 'DS O/T', 'Coaster Normal', 'Coaster O/T'],
    [
        ['1', 'Mooring/unmooring (min $100/$25.90)', '2.00', '3.00', '0.60', '1.12'],
        ['2', 'Other services per 15 min', '100.00', '120.00', '20.30', '27.00'],
    ]
)

# --- CLAUSE 8 ---
add_heading_styled(doc, 'Clause 8: Supply of Fresh Water to Vessels', 2)
add_body(doc, 'Rate per ton (1,000 litres) or HTN (USD):')
add_branded_table(doc,
    ['#', 'Service', 'Rate (USD)'],
    [
        ['1', 'Shore hydrants direct to ship', '4.00'],
        ['2', 'In stream (min $95)', '8.00'],
        ['3(a)', 'Outer anchorage (min $145)', '14.00'],
        ['3(b)', 'From barges', '8.00'],
        ['3(c)', 'From water bowsers', '14.00'],
        ['4', 'Metre testing fee', '10.00'],
    ],
    col_widths=[0.5, 4.0, 1.2]
)
add_body(doc, 'Detention: $26/half-hr (2215-0700 & weekends), $13/half-hr (0700-2300). Traditional vessels: $0.20/GRT.')

# --- CLAUSE 9 ---
add_heading_styled(doc, 'Clause 9: Garbage Disposal', 2)
add_branded_table(doc,
    ['Service', 'Rate (USD)'],
    [
        ['Per receptacle per day', '13.00'],
        ['Vehicle hire per trip (regular)', '114.20'],
        ['Vehicle hire per trip (overtime)', '125.00'],
    ]
)

# --- CLAUSE 11 ---
add_heading_styled(doc, 'Clause 11: Hire of Staff and Labour', 2)
add_body(doc, 'Rate per man per hour or part thereof (USD):')
add_branded_table(doc,
    ['#', 'Category', 'Regular', 'Overtime'],
    [
        ['3(a)', 'Labourer', '2.00', '4.00'],
        ['3(b)', 'Watchman/Security Guard/Sorter', '2.00', '4.00'],
        ['3(c)', 'Crane/Winch Operator', '2.50', '5.00'],
        ['3(d)', 'Forklift Operator', '2.50', '5.00'],
        ['3(e)', 'Serang', '2.50', '5.00'],
        ['3(f)', 'Clerk', '2.50', '5.00'],
        ['3(g)', 'Foreman', '3.00', '6.00'],
        ['3(h)', 'Asst. Operations Officer', '3.50', '7.00'],
        ['3(i)', 'Operations Officer', '4.50', '9.00'],
        ['4(a)', 'Crane/Winch Foreman', '3.50', '7.00'],
        ['5(a)', 'Principal Fire & Safety Officer', '5.00', '10.00'],
        ['5(b)', 'Senior Fire & Safety Officer', '4.50', '9.00'],
        ['5(c)', 'Fire & Safety Officer', '4.00', '8.00'],
        ['5(d)', 'Fire & Safety Inspector', '3.00', '6.00'],
        ['5(e)', 'Leading Fireman', '2.50', '5.00'],
        ['5(g)', 'Firefighter', '2.00', '4.00'],
    ],
    col_widths=[0.5, 3.0, 1.0, 1.0]
)

add_page_break(doc)

# --- CLAUSE 12 ---
add_heading_styled(doc, 'Clause 12: Hire of Equipment', 2)
add_body(doc, 'Rate per hour or part thereof (USD):')
add_branded_table(doc,
    ['#', 'Equipment', 'Within Port', 'Outside Port'],
    [
        ['1(a)', 'Forklift up to 5T', '16.00', '32.00'],
        ['1(b)', 'Forklift 5-10T', '20.00', '40.00'],
        ['1(c)', 'Forklift 11-16T', '60.00', '120.00'],
        ['1(d)', 'Forklift 16-60T', '90.00', '180.00'],
        ['2(a)', 'Mobile crane up to 5T', '40.00', '80.00'],
        ['2(b)', 'Mobile crane 5-10T', '60.00', '120.00'],
        ['2(c)', 'Mobile crane 10-20T', '80.00', '160.00'],
        ['2(d)', 'Mobile crane 20-40T', '100.00', '200.00'],
        ['2(e)', 'Mobile crane 40T+', '150.00', '300.00'],
        ['3(a)', 'Quay crane up to 5T', '100.00', '—'],
        ['3(b)', 'Quay crane 5T+', '120.00', '—'],
        ['4(a)(i)', 'Floating crane up to 60T (min 2hr)', '200.00', '400.00'],
        ['4(a)(ii)', 'Floating crane 60T+', '300.00', '600.00'],
        ['5', 'Pilot boat (incl. crew)', '500.00', '—'],
    ],
    col_widths=[0.6, 3.0, 1.2, 1.2]
)

add_body(doc, 'Miscellaneous Equipment (per calendar day):')
add_branded_table(doc,
    ['Equipment', 'Within Port', 'Outside Port'],
    [
        ['Pallets each', '2.00', '4.00'],
        ['Tarpaulins each', '20.00', '40.00'],
        ['Gangways each', '40.00', '80.00'],
        ['Lorry per trip (regular/overtime)', '100.00 / 200.00', '—'],
        ['Port trailer', '100.00', '200.00'],
    ]
)

# --- CLAUSE 14 ---
add_heading_styled(doc, 'Clause 14: Stevedoring', 2)
add_body(doc, 'Rate per Harbour Tonne or part thereof (USD):')
add_branded_table(doc,
    ['#', 'Cargo Type', 'Rate (USD)'],
    [
        ['a(i)', 'Breakbulk cargo (discharge/load/shift/shut-out)', '5.50'],
        ['a(ii)', 'Cargo loaded but not accepted, discharged', '9.00'],
        ['a(iii)', 'Cargo landed & reshipped / hold-to-hold shift', '9.00'],
        ['a(iv)', 'Transshipment cargo', '6.00'],
        ['b', 'Dry bulk (mechanical)', '6.00'],
        ['c(i)', 'Stevedoring + bagging at quay/silo', '6.00'],
        ['c(ii)', 'Bagging charges', '7.50'],
        ['d(i)', 'Passenger/crew baggage (per package)', '2.00'],
        ['d(ii)', 'Mail bags (per bag)', '1.00'],
        ['e(i)', 'Difficult cargo (charcoal, sulphur, DG, etc.)', '7.00'],
        ['e(ii)', 'Loose timber / loose scrap', '8.00'],
        ['e(iii)', 'Cold storage cargo', '12.00'],
        ['f', 'Valuable cargo', '7.00'],
    ],
    col_widths=[0.5, 4.0, 1.0]
)

add_body(doc, 'Heavy Lifts (additional per package):')
add_branded_table(doc,
    ['Weight', 'Rate (USD)'],
    [
        ['5-10 DWT', '10.00'],
        ['10-20 DWT', '15.00'],
        ['20-40 DWT', '25.00'],
        ['40+ DWT', '36.00'],
    ]
)
add_note(doc, 'Overtime: +$500 per gang per shift. DG surcharge: +10%.')

add_page_break(doc)

# --- CLAUSE 27 ---
add_heading_styled(doc, 'Clause 27: Wharfage', 2)
add_branded_table(doc,
    ['#', 'Category', 'Rate (USD)'],
    [
        ['3(a)(i)', 'Imports — Domestic', '1.6% ad valorem'],
        ['3(a)(ii)', 'Imports — Transit (per HTN)', '3.00'],
        ['3(b)(i)', 'Exports — Domestic', '1.0% ad valorem'],
        ['3(b)(ii)', 'Exports — Transit (per HTN)', '3.00'],
        ['3(c)', 'Transshipment (excl. containers)', '0.8% ad valorem'],
        ['3(d)', 'Dhow cargo (per HTN)', '2.00'],
        ['3(e)', 'Fuel/lube oils to vessels (per DWT)', '2.00'],
        ['3(g)', 'Molasses in bulk', '1.0% ad valorem'],
    ],
    col_widths=[0.6, 3.5, 1.5]
)
add_note(doc, 'Ad valorem: Min $200/HTN, Max $2,500/HTN.')

add_body(doc, 'Secret Cargo:')
add_branded_table(doc,
    ['Category', 'Per HTN', 'Per 20ft', 'Per 40ft+'],
    [
        ['Domestic general cargo', '12.00', '—', '—'],
        ['Transit general cargo', '10.00', '—', '—'],
        ['Domestic containers', '—', '250.00', '500.00'],
        ['Transit containers', '—', '200.00', '400.00'],
    ]
)

add_body(doc, 'Containerised Transit Wharfage:')
add_branded_table(doc,
    ['Category', 'Up to 20ft', 'Over 20ft'],
    [
        ['FCL Imports', '90.00', '180.00'],
        ['FCL Exports', '75.00', '150.00'],
    ]
)

# --- CLAUSE 28 ---
add_heading_styled(doc, 'Clause 28: Wayleave Dues', 2)
add_branded_table(doc,
    ['#', 'Category', 'Rate (USD)'],
    [
        ['3(a)', 'Import/Export per HTN (deep sea)', '3.00'],
        ['3(b)', 'Imports containerised — 20ft / 40ft+', '80.00 / 170.00'],
        ['3(c)', 'Exports containerised — 20ft / 40ft+', '65.00 / 140.00'],
    ]
)

# --- CLAUSE 29 ---
add_heading_styled(doc, 'Clause 29: Shorehandling', 2)

add_body(doc, 'Domestic Traffic — Break Bulk (per HTN):')
add_branded_table(doc,
    ['#', 'Service', 'Rate (USD)'],
    [
        ['1(a)(i)', 'Imported dry/breakbulk', '7.00'],
        ['1(a)(ii)', 'Domestic exports', '3.50'],
        ['1(a)(iii)', 'Transshipment / overlanded cargo', '7.00'],
        ['1(a)(iv)', 'Shut-out cargo', '1.50'],
        ['1(a)(v)', 'Transfer within port area', '3.50'],
        ['1(a)(vi)', 'Transfer to Customs warehouse', '3.50'],
        ['1(a)(vii)', 'Direct delivery from vessel to road/rail', '6.00'],
    ],
    col_widths=[0.6, 3.5, 1.0]
)

add_body(doc, 'Domestic Containers (per unit):')
add_branded_table(doc,
    ['Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['FCL Containers', '90.00', '135.00'],
        ['Stripping/Stuffing', '70.00', '140.00'],
        ['Empty Containers', '10.00', '20.00'],
        ['Verification-FCL', '90.00', '140.00'],
        ['Extra movement / lift on-off', '10.00', '20.00'],
        ['Change of status', '80.00', '135.00'],
    ]
)

add_body(doc, 'Transit Traffic — Break Bulk (per HTN):')
add_branded_table(doc,
    ['Service', 'Rate (USD)'],
    [
        ['Imported dry/breakbulk', '6.00'],
        ['Direct delivery to road/rail', '5.00'],
        ['Transit exports', '3.00'],
    ]
)

add_body(doc, 'Transit Containers (per unit):')
add_branded_table(doc,
    ['Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['FCL Containers', '80.00', '120.00'],
        ['Verification-FCL', '80.00', '140.00'],
        ['Empty Container', '10.00', '20.00'],
        ['Extra movement', '10.00', '20.00'],
        ['Stripping/Stuffing', '70.00', '140.00'],
        ['Change of status', '70.00', '135.00'],
    ]
)
add_note(doc, 'Surcharges: DG +10%, Over-dimension +30%, Cold storage +30%.')

add_page_break(doc)

# --- CLAUSE 32 ---
add_heading_styled(doc, 'Clause 32: Storage', 2)

add_body(doc, '1. Domestic Cargo — Break Bulk (per HTN/day):')
add_branded_table(doc,
    ['Period', 'Import Rate', 'Export Rate'],
    [
        ['Free period', '5 days', '5 days'],
        ['Next 30 days per HTN/day', '1.00', '0.50'],
        ['Thereafter per HTN/day', '1.50', '—'],
    ]
)

add_body(doc, '2. Domestic FCL Containers (per day per unit):')
add_branded_table(doc,
    ['Period', 'Import 20ft', 'Import 40ft+', 'Export 20ft', 'Export 40ft+'],
    [
        ['Free period', '5 days', '5 days', '5 days', '5 days'],
        ['Next period', '20.00 (10 days)', '40.00 (10 days)', '16.00', '32.00'],
        ['Thereafter', '40.00', '80.00', '—', '—'],
    ]
)

add_body(doc, '3. Transit Cargo — Break Bulk (per HTN/day):')
add_branded_table(doc,
    ['Period', 'Import Rate', 'Export Rate'],
    [
        ['Free period', '15 days', '21 days'],
        ['Next 30 days per HTN/day', '1.00', '0.50'],
        ['Thereafter per HTN/day', '1.50', '—'],
    ]
)

add_body(doc, '4. Transit FCL Containers (per day per unit):')
add_branded_table(doc,
    ['Period', 'Import 20ft', 'Import 40ft+', 'Export 20ft', 'Export 40ft+'],
    [
        ['Free period', '15 days', '15 days', '21 days', '21 days'],
        ['Next period', '20.00 (6 days)', '40.00 (6 days)', '16.00', '32.00'],
        ['Thereafter', '40.00', '80.00', '—', '—'],
    ]
)

add_body(doc, '5. Empty Containers (per day per unit):')
add_branded_table(doc,
    ['Period', 'Up to 20ft', 'Over 20ft'],
    [
        ['First 5 days', 'Free', 'Free'],
        ['Next 10 days', '4.00', '8.00'],
        ['Thereafter', '8.00', '16.00'],
        ['From outside port: 3 days free', '—', '—'],
    ]
)
add_note(doc, 'DG Storage: After 24 hours free, +20% surcharge on applicable rates.')
add_note(doc, 'Transshipment loose: 10 days free, then $0.50/HTN/day.')
add_note(doc, 'Coastwise/Dhow: 3 days free, then $0.30/tonne/day.')

add_page_break(doc)

# --- CLAUSE 33 ---
add_heading_styled(doc, 'Clause 33: Coastwise Cargo', 2)
add_body(doc, 'Break Bulk — Domestic (per HTN):')
add_branded_table(doc,
    ['Service', 'Rate (USD)'],
    [
        ['Wharfage', '2.00'],
        ['Shorehandling', '2.00'],
        ['Stevedoring', '2.00'],
        ['Storage per day (3 days free)', '0.30'],
    ]
)
add_body(doc, 'Containerised — Domestic (per unit):')
add_branded_table(doc,
    ['Service', 'Stuffed 20ft', 'Stuffed 40ft+', 'Empty 20ft', 'Empty 40ft+'],
    [
        ['Wharfage', '37.50', '75.00', '0.75', '1.50'],
        ['Shorehandling', '37.50', '75.00', '3.00', '6.00'],
        ['Stevedoring', '37.50', '58.25', '3.75', '5.63'],
        ['Shifting', '22.50', '22.50', '—', '—'],
    ]
)

# --- CLAUSE 36 ---
add_heading_styled(doc, 'Clause 36: Container Handling — DCT Berths 8-11', 2)
add_branded_table(doc,
    ['#', 'Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['(a)', 'Stevedoring FCL', '80.00', '120.00'],
        ['(b)', 'Stevedoring LCL', '160.00', '255.00'],
        ['(c)', 'Stevedoring Empty', '40.00', '60.00'],
        ['(d)', 'Shifting Containers', '100.00', '200.00'],
    ]
)
add_note(doc, 'DG surcharge: +10% on (a),(b),(d). Over-dimension: +30%. Overtime: +$500/gang/shift.')

# --- CLAUSE 37 ---
add_heading_styled(doc, 'Clause 37: Container Handling — Berths 0-7', 2)
add_branded_table(doc,
    ['#', 'Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['(a)', 'Stevedoring FCL', '100.00', '150.00'],
        ['(b)', 'Stevedoring LCL', '170.00', '270.00'],
        ['(c)', 'Stevedoring Empty', '50.00', '70.00'],
        ['(d)', 'Shifting Containers', '115.00', '230.00'],
    ]
)
add_note(doc, 'DG surcharge: +10% on (a),(b),(d). Over-dimension: +30%. Overtime: +$500/gang/shift.')

# --- CLAUSE 38 ---
add_heading_styled(doc, 'Clause 38: Roll On-Roll Off Operations', 2)
add_branded_table(doc,
    ['#', 'Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['(a)', 'Stevedoring FCL', '70.00', '105.00'],
        ['(b)', 'Stevedoring LCL', '140.00', '240.00'],
        ['(c)', 'Stevedoring Empty', '30.00', '40.00'],
        ['(d)', 'Shifting Containers', '80.00', '120.00'],
    ]
)
add_note(doc, 'DG surcharge: +10% on (a),(b),(d). Over-dimension: +30%. Overtime: +$500/gang/shift.')

add_page_break(doc)

# --- CLAUSE 39 ---
add_heading_styled(doc, 'Clause 39: Other Container Service Charges', 2)

add_body(doc, '1. Reefer Containers (per container per day):')
add_branded_table(doc,
    ['Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['Power supply per day', '8.00', '12.00'],
        ['Storage per day (after 48hr free)', '20.00', '40.00'],
    ]
)

add_body(doc, '2. Change of Container Status:')
add_branded_table(doc,
    ['Timing', 'Up to 20ft', 'Over 20ft'],
    [
        ['Notice submitted', '25.00', '35.00'],
        ['48+ hrs before vessel arrival', 'Free', 'Free'],
        ['5+ days after vessel arrival', '50.00 (double)', '70.00 (double)'],
    ]
)

add_body(doc, '3. Stuffing/Stripping: $70.00 (20ft) / $140.00 (40ft+)')

add_body(doc, '4. Transshipment Containers:')
add_branded_table(doc,
    ['Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['Stevedoring (inward + outward)', '90.00', '135.00'],
        ['Storage free period', '15 days', '15 days'],
        ['Storage thereafter per day', '15.00', '30.00'],
        ['Transfer', '10.00', '15.00'],
    ]
)

add_body(doc, '5. Shut-out Containers:')
add_branded_table(doc,
    ['Service', 'Up to 20ft', 'Over 20ft'],
    [
        ['Shut-out charges', '50.00', '75.00'],
        ['Removal charges', '30.00', '45.00'],
        ['Storage free period', '5 days', '5 days'],
        ['Storage thereafter per day', '15.00', '30.00'],
    ]
)

add_body(doc, '6. Wharfage on Empty Containers: $3.00 (20ft) / $6.00 (40ft+)')

# --- CLAUSE 40 ---
add_heading_styled(doc, 'Clause 40: Handling Charges for Bulk Oils', 2)
add_branded_table(doc,
    ['#', 'Service', 'Rate per HTN (USD)'],
    [
        ['(a)', 'Handling at KOJ and other terminals', '1.50'],
        ['(b)', 'Crude products at SBM', '0.50'],
        ['(c)', 'Other liquid products at SBM (non-crude)', '2.50'],
    ]
)

# --- CLAUSE 41 ---
add_heading_styled(doc, 'Clause 41: Grain Terminal Services', 2)
add_branded_table(doc,
    ['#', 'Service', 'Domestic', 'Transit'],
    [
        ['1(a)', 'Shore Handling', '7.00', '6.00'],
        ['1(b)', 'Bagging bulk at silo', '8.00', '8.00'],
        ['1(c)', 'Intake / out-take', '4.00', '4.00'],
        ['1(d)', 'Recirculation', '1.00', '1.00'],
        ['1(e)', 'Aeration', '0.50', '0.50'],
        ['1(f)', 'Fumigation', '1.00', '1.00'],
        ['1(g)', 'Tipping truck hire', '2.00', '2.00'],
    ],
    col_widths=[0.5, 3.0, 1.0, 1.0]
)
add_body(doc, 'Import Storage: Domestic 10 days free / Transit 15 days free, then $1.00/HTN/day for 30 days, then $1.50/HTN/day.')
add_body(doc, 'Export Storage: 5,000T cushion free (if nominated, shipper 10,000+ tonnes). Excess: $0.50/HTN/day domestic, $0.15 transit.')

add_page_break(doc)

# --- CLAUSE 42 ---
add_heading_styled(doc, 'Clause 42: Special Rate for ICDs / Dry Ports', 2)

add_body(doc, 'Package Rates (all services as a bundle):')
add_branded_table(doc,
    ['ICD', 'Services Included', 'Up to 20ft', 'Over 20ft'],
    [
        ['Ex NASACO Yard', '(a)-(f): Lift empty, offload, place in shed, stuff/strip, lift full, transport to port', '280.00', '510.00'],
        ['Ubungo ICD', '(a)-(e): Same minus transport (conveyance separate)', '180.00', '310.00'],
        ['Other ICDs (Kwala Ruvu etc.)', '(a)-(e): Same as Ubungo (conveyance separate)', '180.00', '310.00'],
    ],
    col_widths=[1.5, 2.5, 0.9, 0.9]
)
add_note(doc, 'If not all services used as a package, individual clause rates apply.')

add_body(doc, 'ICD Storage:')
add_branded_table(doc,
    ['Location', 'Rule'],
    [
        ['Ex NASACO — exports', 'Per Clause 32'],
        ['Ubungo ICD — exports', '3,500T cushion free, then Clause 32'],
        ['Ubungo ICD — imports', 'Per Clause 32'],
        ['Other ICDs — exports', '5,000T cushion free, then Clause 32'],
    ]
)

add_body(doc, 'Other ICDs — Imports (Port Extension Mode):')
add_branded_table(doc,
    ['Period', 'Up to 20ft', 'Over 20ft'],
    [
        ['Local: first 30 days', 'Free', 'Free'],
        ['Transit: first 60 days', 'Free', 'Free'],
        ['Local: next 14 days (31-44)', '20.00', '40.00'],
        ['Transit: next 14 days (61-74)', '20.00', '40.00'],
        ['Thereafter (both)', '40.00', '80.00'],
    ]
)

add_body(doc, 'General Cargo at ICDs: Local 30 days free / Transit 60 days free, then $0.50/HTN/day.')
add_body(doc, 'DG at all ICDs: 24 hours free, then Clause 32 + 20% surcharge.')
add_note(doc, 'Port Extension Mode imports: conveyance borne by TPA. Overstayed/shipper nomination: on consignee account.')

add_page_break(doc)

# ======================================================================
# PART IV: CLAUSE 43
# ======================================================================
add_heading_styled(doc, 'PART IV: CLAUSE 43 — Miscellaneous Provisions', 1)
add_gold_divider(doc)

add_heading_styled(doc, 'Weighing and Measuring', 2)
add_branded_table(doc,
    ['#', 'Service', 'Rate (USD)'],
    [
        ['1(a)', 'Weighing packages singly, per 50 kg', '1.00'],
        ['1(b)', 'Weighing collectively, per tonne', '4.00'],
        ['2(a)', 'VGM — Dar es Salaam (per container)', '60.00'],
        ['2(b)', 'VGM — Tanga / Mtwara (per container)', '30.00'],
        ['3(a)', 'Measuring imports/exports per tonne or CBM', '4.00'],
        ['3(b)', 'Measuring baggage per piece', '1.00'],
    ],
    col_widths=[0.5, 3.5, 1.2]
)

add_heading_styled(doc, 'Removal, Sorting, and Other Services', 2)
add_branded_table(doc,
    ['Service', 'Rate per HTN (USD)'],
    [
        ['Removal of exports — DSM', '4.00'],
        ['Removal — Mtwara/Tanga (sender labour)', '2.00'],
        ['Removal — Mtwara/Tanga (port labour)', '3.00'],
        ['Sorting exports', '2.00'],
        ['Sorting imports — sorting only', '3.00'],
        ['Sorting imports — removal to another shed', '4.00'],
        ['Bagging — manual (per bag)', '2.00'],
        ['Bagging — mechanised (per bag)', '1.00'],
        ['Patching/re-sewing/banding (per unit)', '1.00'],
        ['Double handling cereals (per DWT)', '3.00'],
        ['Fumigation — rice (per DWT)', '4.00'],
        ['Fumigation — other (per package)', '1.00'],
    ]
)

add_page_break(doc)

# ======================================================================
# APPENDIX A: FREE STORAGE PERIODS
# ======================================================================
add_heading_styled(doc, 'Appendix A: Free Storage Periods Summary', 1)
add_gold_divider(doc)

add_branded_table(doc,
    ['Cargo Type', 'Domestic', 'Transit'],
    [
        ['Break bulk import', '5 days', '15 days'],
        ['Break bulk export', '5 days', '21 days'],
        ['FCL container import', '5 days', '15 days'],
        ['FCL container export', '5 days', '21 days'],
        ['Empty container (from ship)', '5 days', '5 days'],
        ['Empty container (from outside port)', '3 days', '3 days'],
        ['Transshipment (loose)', '10 days', '10 days'],
        ['Transshipment (container, Clause 39)', '15 days', '15 days'],
        ['Shut-out container', '5 days', '5 days'],
        ['Coastwise / Dhow cargo', '3 days', '—'],
        ['Grain terminal import', '10 days', '15 days'],
        ['Reefer storage', '48 hours', '48 hours'],
        ['DG storage (all locations)', '24 hours', '24 hours'],
        ['ICD import — local (port extension)', '30 days', '—'],
        ['ICD import — transit (port extension)', '—', '60 days'],
        ['Tanga/Mtwara transshipment containers', '30 days', '30 days'],
    ],
    col_widths=[3.5, 1.2, 1.2]
)

doc.add_paragraph()

# ======================================================================
# APPENDIX B: SURCHARGE RULES
# ======================================================================
add_heading_styled(doc, 'Appendix B: Global Surcharge Rules Summary', 1)
add_gold_divider(doc)

add_branded_table(doc,
    ['Surcharge', 'Percentage', 'Applicable Clauses'],
    [
        ['Dangerous Goods — Stevedoring', '+10%', 'Clauses 14, 36, 37, 38'],
        ['Dangerous Goods — Shorehandling', '+10%', 'Clause 29'],
        ['Dangerous Goods — Storage', '+20%', 'Clause 32 (after 24hr free)'],
        ['Over-dimension Containers', '+30%', 'Clauses 29, 36, 37, 38'],
        ['Cold Storage Cargo — Handling', '+30%', 'Clause 29'],
        ['Overtime Gang Charge', '+$500/gang/shift', 'Clauses 14, 36, 37, 38'],
    ],
    col_widths=[2.5, 1.5, 2.0]
)

doc.add_paragraph()
doc.add_paragraph()

# Final note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('End of Document')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = BRAND_PURPLE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tanzania Ports Authority  |  P.O. Box 9184, Dar es Salaam  |  www.ports.go.tz')
run.font.size = Pt(8)
run.font.color.rgb = MID_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by 7Square Inc.  |  www.7squareinc.com')
run.font.size = Pt(8)
run.font.color.rgb = BRAND_GOLD

# ======================================================================
# SAVE
# ======================================================================
doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
