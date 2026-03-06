"""
Generate branded 7Square SOP Word document from SOP-7SQUARE-ICD-PLATFORM.md
Matches the markdown 100% in content, with 7Square branding (purple/gold).
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# === Brand colors ===
BRAND_PURPLE = RGBColor(0x1E, 0x0A, 0x3C)
BRAND_GOLD = RGBColor(0xC5, 0xA4, 0x4E)
BRAND_LIGHT_PURPLE = RGBColor(0x3D, 0x1D, 0x72)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
CODE_BG = "F0F0F0"

BASE = r"c:\Users\ernes\Documents\Projects\tpa_sea_ports"
LOGO = os.path.join(BASE, "The _7Square full color logo.png")
OUTPUT = os.path.join(BASE, "7Square_SOP_ICD_Platform.docx")

doc = Document()

# === Page setup ===
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = DARK_GRAY

# === Helpers ===

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_branded_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        set_cell_shading(cell, "1E0A3C")
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            if val.startswith('**') and val.endswith('**'):
                run = p.add_run(val[2:-2])
                run.bold = True
            else:
                run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8F6FB")
    doc.add_paragraph()

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = BRAND_PURPLE
        # Add gold underline bar
        bar = doc.add_paragraph()
        bar_run = bar.add_run("_" * 80)
        bar_run.font.color.rgb = BRAND_GOLD
        bar_run.font.size = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = BRAND_LIGHT_PURPLE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = BRAND_PURPLE

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    # Parse inline bold markers **text**
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        if i % 2 == 1:  # odd segments are bold
            run.bold = True
        if bold:
            run.bold = True
        if italic:
            run.italic = True
    return p

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    # Parse inline bold/code
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        # Handle backtick code within
        code_parts = part.split('`')
        for j, cp in enumerate(code_parts):
            if not cp:
                continue
            run = p.add_run(cp)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 1:
                run.bold = True
            if j % 2 == 1:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_GRAY
        # Apply background shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        pPr.append(shading)

def add_separator():
    p = doc.add_paragraph()
    run = p.add_run("_" * 80)
    run.font.color.rgb = BRAND_GOLD
    run.font.size = Pt(6)


# =====================================================================
# COVER PAGE
# =====================================================================

# Logo
if os.path.exists(LOGO):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO, width=Inches(2.5))

# Title
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Standard Operating Procedure (SOP)")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = BRAND_PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("7Square ICD Integrated Operations Platform")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = BRAND_LIGHT_PURPLE
run.font.name = 'Calibri'

for _ in range(2):
    doc.add_paragraph()

# Document info box
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Document No:", "7SQ-SOP-001"),
    ("Effective Date:", "March 2026"),
    ("Prepared for:", "7Square Inland Container Depot, Dar es Salaam"),
    ("Contact:", "www.7squareinc.com | info@7squareinc.com"),
]
for i, (label, value) in enumerate(info_data):
    cell_l = info_table.rows[i].cells[0]
    cell_l.text = ''
    run = cell_l.paragraphs[0].add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = BRAND_PURPLE
    cell_r = info_table.rows[i].cells[1]
    cell_r.text = ''
    run = cell_r.paragraphs[0].add_run(value)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GRAY

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS
# =====================================================================

add_heading("Table of Contents", 1)

toc_items = [
    "1. System Overview",
    "2. Login & User Roles",
    "3. Dashboard",
    "4. Customer Management",
    "5. Vessel Registry",
    "6. Container Management",
    "7. Cargo Registration",
    "8. Warehouse Management",
    "9. Dangerous Goods (IMDG)",
    "10. Reefer Monitoring",
    "11. Document Management",
    "12. Tariff Book Reference",
    "13. Billing & Invoicing",
    "14. Reports & KPIs",
    "15. Common Workflows",
    "16. Appendix: TPA Tariff Quick Reference",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = BRAND_LIGHT_PURPLE

doc.add_page_break()

# =====================================================================
# SECTION 1: System Overview
# =====================================================================

add_heading("1. System Overview", 1)

add_para("The 7Square ICD Platform manages the full lifecycle of container and cargo operations at the Inland Container Depot, including:")

add_bullet("Container receiving, storage, and release")
add_bullet("Bonded and free warehouse management")
add_bullet("Billing based on the **TPA Sea Ports Tariff Book (Feb 2024)** -- all 43 clauses pre-loaded")
add_bullet("Dangerous goods classification per **IMDG Code**")
add_bullet("Reefer container power and temperature monitoring")
add_bullet("Trade document tracking with **15-minute SLA** per TPA Charter")
add_bullet("SOLAS **VGM (Verified Gross Mass)** compliance")
add_bullet("Transit cargo routing to landlocked countries (Zambia, DRC, Burundi, Rwanda, Malawi, Uganda, Zimbabwe)")

add_para("**Tech stack:** Next.js 14, PostgreSQL, Prisma ORM, NextAuth.js")

add_separator()

# =====================================================================
# SECTION 2: Login & User Roles
# =====================================================================

add_heading("2. Login & User Roles", 1)

add_heading("2.1 How to Log In", 3)
add_para("1. Open the platform in your browser: http://localhost:3002/login")
add_para("2. Enter your **email** and **password**")
add_para("3. Click **Sign In**")

add_heading("2.2 Default Accounts (change passwords after first login)", 3)

add_branded_table(
    ["Role", "Email", "Password", "Permissions"],
    [
        ["Admin", "admin@7square.co.tz", "admin123", "Full access to all modules"],
        ["Operations", "ops@7square.co.tz", "admin123", "Containers, cargo, vessels, warehouse"],
        ["Warehouse", "warehouse@7square.co.tz", "admin123", "Warehouse, storage, reefer"],
        ["Billing", "billing@7square.co.tz", "admin123", "Invoicing, tariffs, customers"],
    ]
)

add_heading("2.3 Navigation", 3)
add_para("The left sidebar provides access to all 12 modules:")
add_bullet("Dashboard, Containers, Cargo, Warehouse, Vessels, Tariffs, Billing, Documents, Dangerous Goods, Reefer, Customers, Reports")
add_para("Your name and role appear at the bottom of the sidebar. Click **Sign Out** to log out.")

add_separator()

# =====================================================================
# SECTION 3: Dashboard
# =====================================================================

add_heading("3. Dashboard", 1)
add_para("**Path:** Sidebar > Dashboard")
add_para("The dashboard shows real-time KPI cards:")

add_branded_table(
    ["Card", "What it shows"],
    [
        ["Active Containers", "Total containers currently registered"],
        ["Customers", "Total active customer accounts"],
        ["Warehouses", "Number of operational warehouses"],
        ["Invoices", "Total invoices generated"],
        ["TPA Tariff Clauses", "Number of tariff clauses loaded (should be 43)"],
        ["Tariff Rates Loaded", "Total individual rate items (should be 225)"],
    ]
)
add_para("Use this as your daily operational snapshot.")

add_separator()

# =====================================================================
# SECTION 4: Customer Management
# =====================================================================

add_heading("4. Customer Management", 1)
add_para("**Path:** Sidebar > Customers")

add_heading("4.1 View Customers", 3)
add_bullet("The customer list shows all active customers with their type, country, and counts of containers/invoices.")

add_heading("4.2 Add a New Customer", 3)
add_para("1. Click **\"+ Add Customer\"** button (top right)")
add_para("2. Fill in the form:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Name", "Yes", "Contact person or company name"],
        ["Type", "Yes", "Importer, Exporter, Shipping Agent, Clearing Agent, Ship Owner, Terminal Operator, Transit Client"],
        ["Country", "Yes", "TZ (Tanzania), ZM (Zambia), CD (DRC), BI (Burundi), RW (Rwanda), MW (Malawi), UG (Uganda), ZW (Zimbabwe)"],
        ["Company Name", "No", "Registered company name"],
        ["Email", "No", "Contact email"],
        ["Phone", "No", "Contact phone"],
        ["TIN (Tax ID)", "No", "Tanzania Revenue Authority TIN"],
    ]
)

add_para("3. Click **\"Add Customer\"**")
add_para("4. You will be redirected to the customer list")
add_para("**Important:** For transit clients, always select the correct destination country -- this affects storage free periods (transit gets 15-21 days vs domestic 5 days).")

add_separator()

# =====================================================================
# SECTION 5: Vessel Registry
# =====================================================================

add_heading("5. Vessel Registry", 1)
add_para("**Path:** Sidebar > Vessels")

add_heading("5.1 Register a Vessel", 3)
add_para("1. Click **\"+ Register Vessel\"**")
add_para("2. Fill in the form:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Vessel Name", "Yes", 'e.g. "MSC FLAMINIA"'],
        ["IMO Number", "No", "7-digit IMO number (must be unique)"],
        ["Vessel Type", "Yes", "Container Ship, Bulk Carrier, Tanker, RORO, General Cargo, Dhow, Coaster, Traditional, Tug, Other"],
        ["GRT", "Yes", "Gross Registered Tonnage -- critical for charge calculation (pilotage, port dues, tug services are per 100 GRT)"],
        ["DWT", "No", "Deadweight tonnage"],
        ["LOA", "No", "Length overall in metres"],
        ["Flag State", "No", "e.g. Panama, Liberia, Tanzania"],
        ["Coaster", "Checkbox", "Tick if vessel qualifies for reduced coaster rates"],
    ]
)

add_para("3. Click **\"Register Vessel\"**")
add_para("**Why GRT matters:** Pilotage = $5.50/100 GRT, Port Dues = $13.40/100 GRT, Tug = $14.00/100 GRT. A 50,000 GRT vessel entering port costs approximately: Pilotage $2,750 + Port Dues $6,700 + Tug $7,000 = **$16,450** in vessel charges alone.")

add_separator()

# =====================================================================
# SECTION 6: Container Management
# =====================================================================

add_heading("6. Container Management", 1)
add_para("**Path:** Sidebar > Containers")

add_heading("6.1 View Containers", 3)
add_bullet("Shows all containers with status badges (Arriving, Received, In Storage, Under Operation, Released, Departed)")
add_bullet("VGM and DG indicators shown where applicable")

add_heading("6.2 Register a Container", 3)
add_para("1. Click **\"+ Register Container\"**")
add_para("2. Fill in:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Container Number", "Yes", "ISO 6346 format, e.g. MSCU1234567 (auto-uppercased)"],
        ["Size", "Yes", "20ft, 40ft, or 45ft"],
        ["Type", "Yes", "Dry, Reefer, Open Top, Flat Rack, Tank, Other"],
        ["Customer", "No", "Link to existing customer"],
        ["Seal Number", "No", "Customs/shipping seal"],
        ["Tare Weight", "No", "Empty weight in kg"],
        ["VGM Weight", "No", "SOLAS Verified Gross Mass in kg"],
        ["FCL", "Checkbox", "Full Container Load (default: checked)"],
        ["Empty", "Checkbox", "Empty container"],
        ["Over-dimension", "Checkbox", "Triggers +30% surcharge on handling"],
        ["VGM Certified", "Checkbox", "SOLAS VGM compliance confirmed"],
    ]
)

add_para("3. Click **\"Register Container\"**")

add_heading("6.3 Container Status Lifecycle", 3)
add_code_block(["ARRIVING --> RECEIVED --> IN_STORAGE --> UNDER_OPERATION --> READY_FOR_RELEASE --> RELEASED --> DEPARTED"])

add_para("Status updates happen:")
add_bullet("**RECEIVED:** When container physically arrives at ICD")
add_bullet("**IN_STORAGE:** Automatically set when a storage booking is created")
add_bullet("**UNDER_OPERATION:** During stuffing, stripping, palletising, etc.")
add_bullet("**RELEASED:** When all charges paid and customs clearance obtained")
add_bullet("**DEPARTED:** When container physically leaves the ICD")

add_heading("6.4 Surcharge Flags", 3)

add_branded_table(
    ["Flag", "Surcharge", "Applies To"],
    [
        ["Over-dimension", "+30%", "Shorehandling (Clause 29), Stevedoring (Clauses 36-38)"],
        ["DG (set via Cargo)", "+10% handling, +20% storage", "Clauses 14, 29, 32, 36-38"],
        ["Reefer", "Power: $8/day (20ft), $12/day (40ft)", "Clause 39"],
    ]
)

add_separator()

# =====================================================================
# SECTION 7: Cargo Registration
# =====================================================================

add_heading("7. Cargo Registration", 1)
add_para("**Path:** Sidebar > Cargo")

add_heading("7.1 Register Cargo", 3)
add_para("1. Click **\"+ Register Cargo\"**")
add_para("2. Fill in:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Description", "Yes", 'e.g. "Electronic equipment, cement bags"'],
        ["Cargo Type", "Yes", "Domestic Import, Domestic Export, Transit Import, Transit Export, Transshipment, Coastwise"],
        ["HS Code", "No", "Harmonized System code"],
        ["Weight (kg)", "Yes", "Gross weight"],
        ["Volume (CBM)", "No", "Cubic metres"],
        ["CIF Value (USD)", "No", "For ad valorem wharfage calculation"],
        ["Packages", "No", "Number of packages"],
        ["Customer", "No", "Link to customer"],
        ["Destination Country", "Conditional", "Appears for Transit/Transshipment types"],
        ["Dangerous Goods", "Checkbox", "+10% handling, +20% storage surcharge"],
        ["Cold Storage", "Checkbox", "+30% handling surcharge"],
        ["Valuable Cargo", "Checkbox", "Higher stevedoring rate ($7/HTN vs $5.50/HTN)"],
    ]
)

add_heading("7.2 Harbour Tonne Calculation", 3)
add_para("The system **automatically calculates** harbour tonnes using the formula:")
add_code_block(["Harbour Tonnes (HTN) = MAX(Weight in kg / 1000, Volume in CBM)"])
add_para("This is displayed live on the form as you type. HTN is the billing unit for most break-bulk charges.")
add_para("**Example:** 15,000 kg cargo at 22.5 CBM = MAX(15, 22.5) = **22.5 HTN**")

add_heading("7.3 Cargo Type Impact on Billing", 3)

add_branded_table(
    ["Cargo Type", "Free Storage Period", "Wharfage Basis"],
    [
        ["Domestic Import", "5 days", "1.6% ad valorem"],
        ["Domestic Export", "5 days", "1.0% ad valorem"],
        ["Transit Import", "15 days", "$3/HTN flat"],
        ["Transit Export", "21 days", "$3/HTN flat"],
        ["Transshipment", "10 days", "0.8% ad valorem"],
        ["Coastwise", "3 days", "$2/HTN"],
    ]
)

add_separator()

# =====================================================================
# SECTION 8: Warehouse Management
# =====================================================================

add_heading("8. Warehouse Management", 1)
add_para("**Path:** Sidebar > Warehouse")

add_heading("8.1 View Warehouses", 3)
add_bullet("Shows warehouse cards with type, capacity (TEU), and occupancy bar")

add_heading("8.2 Add a Warehouse", 3)
add_para("1. Click **\"+ Add Warehouse\"**")
add_para("2. Fill in:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Warehouse Name", "Yes", 'e.g. "Bonded Warehouse A"'],
        ["Type", "Yes", "Bonded, Free, Reefer Yard, DG Zone, Open Yard"],
        ["Total Capacity (TEU)", "Yes", "Maximum TEU capacity"],
        ["Location", "No", 'e.g. "Block A, Zone 3"'],
    ]
)
add_para("3. Click **\"Add Warehouse\"**")

add_heading("8.3 Warehouse Types", 3)

add_branded_table(
    ["Type", "Purpose", "Special Rules"],
    [
        ["**Bonded**", "Customs-controlled goods awaiting clearance", "Goods under customs bond"],
        ["**Free**", "Cleared goods, general storage", "Standard rates"],
        ["**Reefer Yard**", "Refrigerated containers", "Power supply $8-12/day, 48hr free storage"],
        ["**DG Zone**", "Dangerous goods (IMDG compliant)", "24hr free, then +20% storage surcharge"],
        ["**Open Yard**", "Break bulk, over-dimension cargo", "Standard rates"],
    ]
)

add_separator()

# =====================================================================
# SECTION 9: Dangerous Goods (IMDG)
# =====================================================================

add_heading("9. Dangerous Goods (IMDG)", 1)
add_para("**Path:** Sidebar > Dangerous Goods")

add_heading("9.1 Register Dangerous Goods", 3)
add_para("1. Click **\"+ Register DG\"**")
add_para("2. Fill in:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["UN Number", "Yes", "e.g. UN1203 (auto-uppercased)"],
        ["IMDG Class", "Yes", "Class 1-9 per IMDG Code (see table below)"],
        ["Proper Shipping Name", "Yes", 'e.g. "METHANOL"'],
        ["Packing Group", "No", "I (Great Danger), II (Medium), III (Minor)"],
        ["Flash Point", "No", 'e.g. "11 C"'],
        ["Container", "No", "Link to container carrying the DG"],
        ["Segregation Group", "No", "For IMDG segregation compliance"],
        ["Notes", "No", "Handling instructions"],
    ]
)
add_para("3. Click **\"Register DG\"**")

add_heading("9.2 IMDG Classes", 3)

add_branded_table(
    ["Class", "Description", "Examples"],
    [
        ["1", "Explosives", "Fireworks, ammunition"],
        ["2.1", "Flammable Gases", "LPG, propane"],
        ["2.2", "Non-Flammable Gases", "Nitrogen, CO2"],
        ["2.3", "Toxic Gases", "Chlorine, ammonia"],
        ["3", "Flammable Liquids", "Methanol, petrol, paint"],
        ["4.1", "Flammable Solids", "Matches, sulphur"],
        ["4.2", "Spontaneously Combustible", "White phosphorus"],
        ["4.3", "Dangerous When Wet", "Sodium, calcium"],
        ["5.1", "Oxidizing Substances", "Ammonium nitrate"],
        ["5.2", "Organic Peroxides", "Methyl ethyl ketone peroxide"],
        ["6.1", "Toxic Substances", "Pesticides, cyanide"],
        ["6.2", "Infectious Substances", "Medical waste"],
        ["7", "Radioactive Material", "Uranium, medical isotopes"],
        ["8", "Corrosives", "Battery acid, bleach"],
        ["9", "Miscellaneous", "Lithium batteries, dry ice"],
    ]
)

add_heading("9.3 DG Surcharge Summary", 3)

add_branded_table(
    ["Charge Type", "Surcharge", "Reference"],
    [
        ["Stevedoring", "+10%", "Clauses 14, 36, 37, 38"],
        ["Shorehandling", "+10%", "Clause 29"],
        ["Storage", "+20% (after 24hr free)", "Clause 32"],
        ["Lighter hire", "Treble (3x) rates", "Clause 6"],
    ]
)

add_separator()

# =====================================================================
# SECTION 10: Reefer Monitoring
# =====================================================================

add_heading("10. Reefer Monitoring", 1)
add_para("**Path:** Sidebar > Reefer")
add_para("This module tracks refrigerated containers connected to power supply at the ICD.")

add_heading("10.1 Key Rates (Clause 39)", 3)

add_branded_table(
    ["Service", "20ft", "40ft+"],
    [
        ["Power supply per day", "$8.00", "$12.00"],
        ["Reefer storage (after 48hr free) per day", "$20.00", "$40.00"],
    ]
)

add_heading("10.2 What to Monitor", 3)
add_bullet("**Set Temperature vs Actual Temperature** -- alerts if deviation exceeds threshold")
add_bullet("**Power Status** -- connected/disconnected")
add_bullet("**Days on Power** -- for billing calculation")

add_separator()

# =====================================================================
# SECTION 11: Document Management
# =====================================================================

add_heading("11. Document Management", 1)
add_para("**Path:** Sidebar > Documents")

add_heading("11.1 Upload a Document", 3)
add_para("1. Click **\"+ Upload Document\"**")
add_para("2. Fill in:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Document Type", "Yes", "See list below"],
        ["Document Number", "No", "e.g. BL-2024-0891"],
        ["File Name", "Yes", "e.g. BL_MSCU1234567.pdf"],
        ["Container", "No", "Link to container"],
        ["Customer", "No", "Link to customer"],
        ["Notes", "No", "Additional notes"],
    ]
)
add_para("3. Click **\"Upload Document\"**")

add_heading("11.2 Document Types", 3)

add_branded_table(
    ["Type", "Abbreviation", "When Used"],
    [
        ["Bill of Lading", "B/L", "Cargo receipt, proof of shipment"],
        ["Delivery Order", "D/O", "Authorization to release cargo"],
        ["Release Order", "R/O", "Customs release authorization"],
        ["Customs Declaration", "-", "Import/export customs filing"],
        ["MSDS", "MSDS", "Material Safety Data Sheet (DG cargo)"],
        ["Shipping Order", "S/O", "Booking confirmation"],
        ["Manifest", "-", "Cargo manifest from vessel"],
        ["VGM Certificate", "VGM", "SOLAS verified gross mass"],
        ["TANCIS Entry", "-", "Tanzania Customs Integrated System entry"],
        ["Packing List", "-", "Itemized cargo contents"],
        ["Commercial Invoice", "C/I", "Seller's invoice for valuation"],
        ["Certificate of Origin", "C/O", "Country of manufacture"],
        ["Fumigation Certificate", "-", "Pest treatment confirmation"],
        ["DG Declaration", "-", "Dangerous goods declaration form"],
        ["Other", "-", "Any other document"],
    ]
)

add_heading("11.3 SLA Target", 3)
add_para("Per TPA Charter: **All documents must be processed within 15 minutes** of receipt. The system tracks processing time against this target.")

add_separator()

# =====================================================================
# SECTION 12: Tariff Book Reference
# =====================================================================

add_heading("12. Tariff Book Reference", 1)
add_para("**Path:** Sidebar > Tariffs")

add_heading("12.1 Browsing Tariffs", 3)
add_para("The tariff page displays all **43 clauses** of the TPA Sea Ports Tariff Book (Feb 2024). Each clause can be expanded to view its rate table.")

add_heading("12.2 Key Clauses for Daily Operations", 3)

add_branded_table(
    ["Clause", "Title", "What It Covers"],
    [
        ["1", "Pilotage Fees", "$5.50/100 GRT per operation"],
        ["2", "Port Dues", "$13.40/100 GRT first 5 days"],
        ["5", "Tug Services", "$14.00/100 GRT per tug"],
        ["14", "Stevedoring", "$5.50/HTN breakbulk"],
        ["27", "Wharfage", "1.6% ad valorem import, $90/$180 FCL transit"],
        ["29", "Shorehandling", "$90/$135 domestic FCL, $80/$120 transit FCL"],
        ["32", "Storage", "Tiered daily rates after free period"],
        ["36", "Container Handling (DCT 8-11)", "$80/$120 stevedoring FCL"],
        ["37", "Container Handling (Berths 0-7)", "$100/$150 stevedoring FCL"],
        ["39", "Other Container Services", "Reefer power, stuffing/stripping"],
        ["42", "ICD Special Rates", "$180/$310 (Ubungo), $280/$510 (NASACO)"],
        ["43", "Miscellaneous", "VGM $60, weighing, sorting"],
    ]
)

add_separator()

# =====================================================================
# SECTION 13: Billing & Invoicing
# =====================================================================

add_heading("13. Billing & Invoicing", 1)
add_para("**Path:** Sidebar > Billing")

add_heading("13.1 View Invoices", 3)
add_bullet("Shows all invoices with status (Draft, Issued, Paid, Overdue)")
add_bullet("Filter by status or customer")

add_heading("13.2 Create a New Invoice", 3)
add_para("1. Click **\"+ New Invoice\"**")
add_para("2. Fill in the header:")

add_branded_table(
    ["Field", "Required", "Notes"],
    [
        ["Customer", "Yes", "Select from registered customers"],
        ["Currency", "Yes", "USD or TZS"],
        ["VAT Rate (%)", "Yes", "Default: 18%"],
        ["Notes", "No", 'e.g. "Container MSCU1234567 services"'],
    ]
)

add_para("3. Add **line items** (click \"+ Add Item\" for more):")

add_branded_table(
    ["Field", "Notes"],
    [
        ["Description", 'Service description, e.g. "Container handling FCL 20ft"'],
        ["Clause Ref", 'TPA clause reference, e.g. "Clause 36"'],
        ["Qty", "Number of units"],
        ["Unit Rate ($)", "Rate per unit from tariff book"],
        ["Line Total", "Auto-calculated: Qty x Unit Rate"],
    ]
)

add_para("4. Review the **totals panel** (bottom right):")
add_bullet("Subtotal (sum of all line items)")
add_bullet("VAT (subtotal x VAT rate)")
add_bullet("**Total** (subtotal + VAT)")
add_para("5. Click **\"Create Invoice (Draft)\"**")

add_heading("13.3 Invoice Number Format", 3)
add_para("Auto-generated: INV-YYYYMMDD-XXXX (e.g. INV-20260306-0001)")

add_heading("13.4 Common Invoice Line Items", 3)
add_para("**For a typical domestic FCL 20ft container (Clause 32 storage + handling):**")

add_branded_table(
    ["Line Item", "Clause", "Rate"],
    [
        ["Container stevedoring FCL 20ft", "Clause 36", "$80.00"],
        ["Shorehandling FCL 20ft domestic", "Clause 29", "$90.00"],
        ["Wharfage FCL 20ft domestic", "Clause 27", "$90.00"],
        ["Storage (days beyond free period x daily rate)", "Clause 32", "$20.00/day"],
    ]
)

add_para("**Add surcharges as separate line items or adjust unit rate:**")
add_bullet("DG container: Add +10% to handling lines, +20% to storage lines")
add_bullet("Over-dimension: Add +30% to handling lines")
add_bullet("Cold storage: Add +30% to shorehandling line")

add_separator()

# =====================================================================
# SECTION 14: Reports & KPIs
# =====================================================================

add_heading("14. Reports & KPIs", 1)
add_para("**Path:** Sidebar > Reports")
add_para("Displays UNCTAD port performance indicators:")

add_branded_table(
    ["KPI", "What It Measures"],
    [
        ["Container Throughput (TEU)", "Total TEU handled per period"],
        ["Average Dwell Time", "Days containers spend in storage"],
        ["Berth Occupancy Rate", "Percentage of time berths are in use"],
        ["Vessel Turnaround Time", "Hours from vessel arrival to departure"],
        ["SLA Compliance Rate", "% of documents processed within 15 min"],
        ["Revenue per TEU", "Average revenue generated per container"],
    ]
)

add_separator()

# =====================================================================
# SECTION 15: Common Workflows
# =====================================================================

add_heading("15. Common Workflows", 1)

add_heading("Workflow A: Receive an Import Container", 2)

add_code_block([
    "Step 1: Ensure the customer exists       --> Customers > + Add Customer (if new)",
    "Step 2: Register the vessel (if new)     --> Vessels > + Register Vessel",
    "Step 3: Register the container           --> Containers > + Register Container",
    "                                             (set size, type, customer, VGM)",
    "Step 4: Register the cargo               --> Cargo > + Register Cargo",
    "                                             (set cargo type, weight, volume)",
    "Step 5: Upload documents                 --> Documents > + Upload Document",
    "                                             (B/L, D/O, customs declaration)",
    "Step 6: If DG cargo                      --> Dangerous Goods > + Register DG",
    "                                             (IMDG class, UN number, MSDS)",
    "Step 7: Container enters storage         --> Status auto-updates to IN_STORAGE",
    "Step 8: Generate invoice                 --> Billing > + New Invoice",
    "                                             (add handling, storage, wharfage lines)",
    "Step 9: Release container                --> Container status > RELEASED",
])

add_heading("Workflow B: Transit Container to Zambia", 2)

add_code_block([
    "Step 1: Register customer (Transit Client, country: ZM)",
    "Step 2: Register container (FCL 40ft)",
    "Step 3: Register cargo (Transit Import, destination: ZM)",
    "                         Weight: 22,000 kg -> HTN auto-calculated",
    "Step 4: Upload B/L + TANCIS entry + customs declaration",
    "Step 5: Note: Transit gets 15 days FREE storage (vs 5 days domestic)",
    "Step 6: Invoice with transit rates:",
    "         - Shorehandling transit FCL 40ft: $120 (Clause 29)",
    "         - Wharfage transit FCL 40ft: $180 (Clause 27)",
    "         - Storage after day 15: $40/day for 40ft (Clause 32)",
    "Step 7: Release for transit trucking to Zambia",
])

add_heading("Workflow C: Dangerous Goods Container", 2)

add_code_block([
    "Step 1: Register container (note: do NOT tick over-dimension unless applicable)",
    "Step 2: Register cargo (tick \"Dangerous Goods\" checkbox)",
    "Step 3: Register DG details (IMDG class, UN number, packing group)",
    "Step 4: Upload MSDS document",
    "Step 5: Assign to DG Zone warehouse",
    "Step 6: DG storage: only 24 hours free, then +20% surcharge",
    "Step 7: Invoice with DG surcharges:",
    "         - Stevedoring: base rate + 10%",
    "         - Shorehandling: base rate + 10%",
    "         - Storage: base rate + 20%",
])

add_heading("Workflow D: Reefer Container", 2)

add_code_block([
    "Step 1: Register container (Type: Reefer, Size: 40ft)",
    "Step 2: Assign to Reefer Yard warehouse",
    "Step 3: Connect to power supply (tracked in Reefer module)",
    "Step 4: Monitor temperature daily",
    "Step 5: Reefer power billing: $12/day for 40ft (Clause 39)",
    "Step 6: Reefer storage: 48 hours free, then $40/day for 40ft",
    "Step 7: Include power + storage charges on invoice",
])

add_separator()

# =====================================================================
# SECTION 16: Appendix
# =====================================================================

add_heading("16. Appendix: TPA Tariff Quick Reference", 1)

add_heading("Storage Free Periods", 3)

add_branded_table(
    ["Cargo Category", "Domestic", "Transit"],
    [
        ["Break bulk import", "5 days", "15 days"],
        ["Break bulk export", "5 days", "21 days"],
        ["FCL container import", "5 days", "15 days"],
        ["FCL container export", "5 days", "21 days"],
        ["Empty container", "5 days", "5 days"],
        ["Transshipment", "10 days", "10 days"],
        ["Coastwise", "3 days", "--"],
        ["ICD local import (port extension)", "30 days", "--"],
        ["ICD transit import (port extension)", "60 days", "--"],
        ["Reefer storage", "48 hours", "48 hours"],
        ["DG storage", "24 hours", "24 hours"],
    ]
)

add_heading("Storage Daily Rates (after free period)", 3)

add_branded_table(
    ["Container", "Tier 1 Rate", "Tier 2 Rate"],
    [
        ["FCL 20ft Import", "$20.00/day (next 10 days)", "$40.00/day (thereafter)"],
        ["FCL 40ft Import", "$40.00/day (next 10 days)", "$80.00/day (thereafter)"],
        ["FCL 20ft Export", "$16.00/day", "--"],
        ["FCL 40ft Export", "$32.00/day", "--"],
        ["Empty 20ft", "$4.00/day (next 10 days)", "$8.00/day (thereafter)"],
        ["Empty 40ft", "$8.00/day (next 10 days)", "$16.00/day (thereafter)"],
    ]
)

add_heading("Surcharge Summary", 3)

add_branded_table(
    ["Surcharge", "Rate", "Applies To"],
    [
        ["Dangerous Goods -- handling", "+10%", "Stevedoring & shorehandling"],
        ["Dangerous Goods -- storage", "+20%", "Storage (after 24hr free)"],
        ["Over-dimension containers", "+30%", "Shorehandling & stevedoring"],
        ["Cold storage cargo", "+30%", "Shorehandling"],
        ["Overtime (3rd shift/weekends)", "Variable", "Labour & equipment"],
    ]
)

add_heading("ICD Package Rates (Clause 42)", 3)

add_branded_table(
    ["ICD Location", "20ft", "40ft"],
    [
        ["NASACO Yard", "$280", "$510"],
        ["Ubungo ICD", "$180", "$310"],
        ["Other ICDs (Kwala Ruvu etc.)", "$180", "$310"],
    ]
)

add_heading("VGM Charges (Clause 43)", 3)

add_branded_table(
    ["Location", "Rate per container"],
    [
        ["Dar es Salaam port", "$60.00"],
        ["Tanga / Mtwara", "$30.00"],
    ]
)

add_separator()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("END OF SOP")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = BRAND_PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("This document should be reviewed and updated whenever TPA tariff rates change or new operational procedures are introduced.")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MID_GRAY
run.font.name = 'Calibri'

# === Save ===
doc.save(OUTPUT)
print(f"SOP document saved to: {OUTPUT}")
